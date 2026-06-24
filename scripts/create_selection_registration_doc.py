from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "/Users/dexley/Documents/TeamFlowAI/output/附件二：选题登记表-已填写.docx"

TITLE = "基于 Spring Boot 与 Vue3 的课程项目小组协作及 AI 项目管理系统设计与实现"
CLASS_VALUE = "请填写"
STUDENT_ID_VALUE = "请填写"
STUDENT_NAME_VALUE = "请填写"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="7F7F7F", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_fixed_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])
                set_cell_margins(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def apply_run_font(run, size=12, bold=False, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph(cell_or_doc, text="", size=12, bold=False, align=None, first_line=False, after=4, line=1.25):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    apply_run_font(r, size=size, bold=bold)
    return p


def replace_first_cell_paragraph(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    apply_run_font(r, size=size, bold=bold)
    return p


def add_page_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    apply_run_font(r, size=16, bold=True)


def add_section_box(doc, heading, paras, min_height_cm=20):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="808080", size="8")
    set_fixed_table_width(table, [15.8])
    cell = table.cell(0, 0)
    cell.height = Cm(min_height_cm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    replace_first_cell_paragraph(cell, heading, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for text in paras:
        paragraph(cell, text, size=12, first_line=True, after=5, line=1.35)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_cover(doc):
    for _ in range(8):
        paragraph(doc, "", size=12, after=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Web应用开发综合实践")
    apply_run_font(r, size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("选题登记表")
    apply_run_font(r, size=22, bold=True)

    table = doc.add_table(rows=3, cols=4)
    set_table_borders(table, color="808080", size="8")
    widths = [2.6, 5.2, 2.6, 5.2]
    set_fixed_table_width(table, widths)
    for row in table.rows:
        row.height = Cm(1.45)
        for cell in row.cells:
            set_cell_margins(cell, top=120, start=160, bottom=120, end=160)

    # Row 1: title spans three value columns.
    replace_first_cell_paragraph(table.cell(0, 0), "题  目", size=12, bold=True)
    merged = table.cell(0, 1).merge(table.cell(0, 3))
    replace_first_cell_paragraph(merged, TITLE, size=12, bold=True)

    replace_first_cell_paragraph(table.cell(1, 0), "班　　级", size=12, bold=True)
    replace_first_cell_paragraph(table.cell(1, 1), CLASS_VALUE, size=12)
    replace_first_cell_paragraph(table.cell(1, 2), "学　号", size=12, bold=True)
    replace_first_cell_paragraph(table.cell(1, 3), STUDENT_ID_VALUE, size=12)

    replace_first_cell_paragraph(table.cell(2, 0), "学生姓名", size=12, bold=True)
    replace_first_cell_paragraph(table.cell(2, 1), STUDENT_NAME_VALUE, size=12)
    replace_first_cell_paragraph(table.cell(2, 2), "", size=12)
    replace_first_cell_paragraph(table.cell(2, 3), "", size=12)

    for c in (table.cell(0, 0), table.cell(1, 0), table.cell(1, 2), table.cell(2, 0)):
        set_cell_shading(c, "F2F2F2")

    doc.add_page_break()


def add_footer_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("第 ")
        apply_run_font(run, size=10)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)
        tail = p.add_run(" 页")
        apply_run_font(tail, size=10)


BACKGROUND = [
    "随着高校实践教学和软件工程课程项目的不断推进，学生小组需要在有限周期内完成需求分析、任务分工、开发协作、过程记录、文档归档和成果汇报等工作。传统方式通常依赖即时通信工具、表格或线下记录，任务状态分散、进度反馈滞后、教师监督成本较高，容易出现职责不清、风险发现不及时、阶段材料缺失等问题。",
    "本课题拟设计并实现 TeamFlowAI 课程项目小组协作与 AI 项目管理系统，面向 Web 应用开发综合实践等课程场景，提供项目管理、成员管理、任务看板、进度记录、文档管理、统计看板和 AI 分析等功能。系统通过前后端分离架构实现业务闭环，并利用大模型接口或本地演示生成逻辑辅助生成周报、风险分析和项目总结。",
    "本选题具有一定的教学实践意义和工程应用价值。一方面，它能够帮助学生以软件工程方式组织课程项目，提升协作效率、过程规范性和项目可视化管理能力；另一方面，它为教师提供项目进展和风险观察入口，便于阶段性指导与验收。通过本系统的设计与实现，能够综合训练 Web 前端开发、后端接口设计、数据库建模、权限认证、数据可视化和 AI 应用集成等能力。"
]


CONTENT = [
    "本课题的研究内容主要围绕课程项目小组协作管理平台展开，系统采用前后端分离模式建设。前端实现登录注册、首页看板、项目列表、项目详情、任务看板、进度记录、文件管理、AI 分析和教师看板等页面；后端提供用户认证、项目管理、成员管理、任务流转、进度记录、文件上传下载、统计分析和 AI 内容生成等 RESTful 接口；数据库负责存储用户、项目、成员、任务、任务日志、周进度、项目文件和 AI 生成记录等核心数据。",
    "系统拟解决的主要问题包括：第一，课程项目成员分工和任务状态缺乏统一管理的问题，通过看板式任务流转记录待开始、进行中、待检查、已完成和阻塞等状态；第二，项目过程材料分散的问题，通过进度记录和文件管理模块沉淀阶段成果；第三，项目进度和风险不易量化的问题，通过统计看板展示参与项目数、任务数、完成率和逾期任务等指标；第四，教师指导和项目验收缺少过程依据的问题，通过教师看板和风险记录辅助监督；第五，项目报告编写效率较低的问题，通过 AI 周报、风险分析和总结报告生成功能提供辅助内容。",
    "本课题最终目标是完成一个可运行、可演示、可扩展的课程项目协作管理系统 MVP。系统需具备基本的角色区分、登录鉴权、核心业务数据增删改查、任务状态流转、统计可视化、文件管理和 AI 辅助分析能力，并能够支撑课程答辩中的完整演示流程。"
]


METHOD = [
    "本课题采用需求分析、系统设计、编码实现、联调测试和演示验证相结合的研究方法。首先分析课程项目协作场景中的用户角色、业务流程和数据对象，确定学生、教师、管理员等角色的主要功能需求；其次进行数据库结构设计和系统模块划分，明确前端页面、后端接口、服务层逻辑和数据表之间的关系；最后基于实际项目代码完成系统实现与功能验证。",
    "技术路线方面，前端采用 Vue 3 与 Vite 搭建单页应用，使用 Vue Router 管理页面路由，使用 Pinia 维护用户状态，使用 Axios 封装接口请求并通过请求拦截器携带 JWT，使用 Element Plus 构建表单、表格、对话框和按钮等交互组件，使用 ECharts 完成任务状态和项目统计的可视化展示。",
    "后端采用 Java 17 与 Spring Boot 3 构建 RESTful 服务，使用 Spring Web 实现接口层，使用 Spring Security 与 JWT 完成无状态登录认证和访问控制，使用 BCrypt 进行密码加密，使用 MyBatis-Plus 简化实体、Mapper 和数据库访问逻辑，使用 MySQL 8 作为关系型数据库，使用 Knife4j/OpenAPI 3 生成接口文档。AI 模块通过可配置的大模型 API 地址、模型名称和 API Key 发起 Chat Completions 格式请求；在未配置真实密钥时，系统启用本地演示生成逻辑，保证离线演示的稳定性。",
    "整体实现流程为：建立数据库并导入初始化脚本；完成后端实体、DTO、VO、Mapper、Service 和 Controller 分层开发；配置安全过滤器、跨域和接口文档；完成前端工程结构、路由守卫、接口封装和页面组件；进行前后端联调，验证登录、项目、任务、进度、文件、看板和 AI 生成流程；最后结合测试账号完成系统演示和问题修正。"
]


REFERENCES = [
    "[1] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
    "[2] 张海藩. 软件工程导论[M]. 北京: 清华大学出版社, 2013.",
    "[3] Craig Walls. Spring in Action[M]. Manning Publications, 2022.",
    "[4] Spring. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/",
    "[5] Spring. Spring Security Reference Documentation[EB/OL]. https://docs.spring.io/spring-security/",
    "[6] MyBatis-Plus. MyBatis-Plus 官方文档[EB/OL]. https://baomidou.com/",
    "[7] Vue.js. Vue 3 官方文档[EB/OL]. https://cn.vuejs.org/",
    "[8] Vite. Vite 官方文档[EB/OL]. https://cn.vitejs.dev/",
    "[9] Element Plus. Element Plus 官方文档[EB/OL]. https://element-plus.org/zh-CN/",
    "[10] Apache ECharts. ECharts 官方文档[EB/OL]. https://echarts.apache.org/zh/",
    "[11] OpenAPI Initiative. OpenAPI Specification[EB/OL]. https://spec.openapis.org/oas/latest.html",
    "[12] DeepSeek. DeepSeek API Documentation[EB/OL]. https://api-docs.deepseek.com/"
]


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_page_title(doc, "Web应用开发综合实践选题登记表")
    add_section_box(doc, "一、选题的背景与意义：", BACKGROUND, min_height_cm=20.7)
    doc.add_page_break()
    add_section_box(doc, "二、研究的基本内容与拟解决的主要问题：", CONTENT, min_height_cm=22.0)
    doc.add_page_break()
    add_section_box(doc, "三、研究的方法与技术路线：", METHOD, min_height_cm=22.0)
    doc.add_page_break()
    add_section_box(doc, "四、主要参考文献：", REFERENCES, min_height_cm=17.2)
    add_footer_page_numbers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
